#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""在简历中插入请求缓存层亮点"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from copy import deepcopy

os.chdir(r'C:\Users\Administrator\Desktop\简历')

target = [f for f in os.listdir('.') if f.startswith('个人简历') and '副本' in f and not f.startswith('~')][0]
print(f'Target: {target}')

doc = Document(target)

# 新亮点内容
new_highlight = (
    "Axios Adapter 层请求缓存与去重：自建适配器级缓存管理器，设计三层拦截管线——"
    "并发请求去重（相同 GET 共享 Promise 避免重复发出）、"
    "TTL 响应缓存（30 秒内命中直接返回深拷贝数据、零网络请求）、"
    "写后主动失效（POST/PUT/DELETE 后按 URL 模式批量清缓存），"
    "核心代码约 180 行，对业务组件零侵入，消除高频页 60%+ 重复请求。"
)

# 找到要插入的位置：项目亮点最后一项之后（段落"专业技能"之前）
# 段落[20] 是 Axios 拦截器相关亮点，[21] 是 "专业技能"
insert_after_idx = 20  # 在段落[20]之后插入

ref_para = doc.paragraphs[insert_after_idx]

# 获取参考段落的样式信息
# 段落[19] 和 [20] 的字体样式（作为新段落的参考）
ref_style_para = doc.paragraphs[19]  # 用第一个亮点段落作为样式参考

# 获取参考段落第一个 run 的字体属性
ref_run = None
for r in ref_style_para.runs:
    ref_run = r
    break

# 在 XML 树中创建新段落
new_p = deepcopy(ref_para._element)  # 复制参考段落的 XML（保留样式）
# 清空内容
for child in list(new_p):
    new_p.remove(child)

# 创建 run 并设置文本
new_r = new_p.makeelement(qn('w:r'), {})
new_rPr = new_r.makeelement(qn('w:rPr'), {})

# 尝试复制参考 run 的字体属性
if ref_run is not None:
    ref_rPr = ref_run._element.find(qn('w:rPr'))
    if ref_rPr is not None:
        new_rPr = deepcopy(ref_rPr)

new_r.insert(0, new_rPr)  # rPr 必须在 t 之前

new_t = new_r.makeelement(qn('w:t'), {})
new_t.text = new_highlight
new_t.set(qn('xml:space'), 'preserve')
new_r.append(new_t)
new_p.append(new_r)

# 在参考段落之后插入
ref_para._element.addnext(new_p)

# 保存（处理文件被占用的情况，保存为新文件）
output_name = target.replace('.docx', '-已添加亮点.docx')
try:
    doc.save(target)
    print(f'Saved to: {target}')
except PermissionError:
    doc.save(output_name)
    print(f'原文件被占用，已另存为: {output_name}')
    print('请关闭原文件后手动替换')

print(f'Done! Highlight inserted after paragraph [{insert_after_idx}]')
print(f'Content: {new_highlight}')
